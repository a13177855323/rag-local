"""
文档处理模块 - 负责各种格式文档的读取、解析和分块
支持 PDF、Word、Markdown、纯文本等多种格式
采用策略模式实现可扩展的文件处理器架构
"""

import os
import re
import hashlib
from typing import List, Dict, Optional, Protocol, runtime_checkable
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from concurrent.futures import ThreadPoolExecutor, as_completed

from PyPDF2 import PdfReader
from docx import Document
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter

from backend.config import settings


@runtime_checkable
class FileProcessor(Protocol):
    """文件处理器协议 - 定义所有文件处理器必须实现的接口"""

    def read(self, file_path: str) -> str:
        """
        读取文件内容

        Args:
            file_path: 文件路径

        Returns:
            str: 提取的文本内容

        Raises:
            Exception: 读取失败时抛出相应异常
        """
        ...

    def supports(self, file_ext: str) -> bool:
        """
        检查是否支持该文件类型

        Args:
            file_ext: 文件扩展名（含小数点，如 '.pdf'）

        Returns:
            bool: 支持返回True，否则返回False
        """
        ...


@dataclass
class DocumentChunk:
    """文档分块数据结构 - 存储单个文本块的元数据和内容"""

    content: str
    chunk_id: int
    total_chunks: int
    filename: str
    file_type: str
    start_index: int = 0
    end_index: int = 0
    content_hash: str = field(init=False)

    def __post_init__(self) -> None:
        """计算内容哈希值用于去重和校验"""
        self.content_hash = hashlib.md5(self.content.encode('utf-8')).hexdigest()

    def to_dict(self) -> Dict:
        """转换为字典格式（用于API响应）"""
        return {
            "id": f"{self.filename}_{self.chunk_id}",
            "content": self.content,
            "metadata": {
                "filename": self.filename,
                "chunk_id": self.chunk_id,
                "total_chunks": self.total_chunks,
                "file_type": self.file_type,
                "start_index": self.start_index,
                "end_index": self.end_index,
                "content_hash": self.content_hash
            }
        }


class PDFProcessor:
    """PDF文件处理器"""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() == '.pdf'

    def read(self, file_path: str) -> str:
        """
        读取PDF文件内容

        优化点：
        1. 添加页码标记
        2. 处理常见的PDF解析问题
        3. 改进文本提取质量
        """
        text_parts = []
        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    # 清理多余的空白字符
                    cleaned_text = self._clean_pdf_text(page_text)
                    text_parts.append(f"[第{page_num}页]\n{cleaned_text}")
                else:
                    print(f"警告：PDF第{page_num}页未提取到有效内容")

            return '\n\n'.join(text_parts)
        except Exception as e:
            raise RuntimeError(f"读取PDF文件失败: {str(e)}") from e

    def _clean_pdf_text(self, text: str) -> str:
        """清理PDF提取的文本，移除多余的空白和换行"""
        # 合并连续的空白字符
        text = re.sub(r'\s+', ' ', text)
        # 修复断行（针对英文）
        text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)
        return text.strip()


class WordProcessor:
    """Word文档处理器"""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in ('.docx', '.doc')

    def read(self, file_path: str) -> str:
        """
        读取Word文档内容

        优化点：
        1. 保留段落结构
        2. 提取表格内容
        3. 处理列表格式
        """
        try:
            doc = Document(file_path)
            content_parts = []

            # 处理段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())

            # 处理表格（简化版）
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_content.append(' | '.join(row_cells))
                if table_content:
                    content_parts.append('[表格]\n' + '\n'.join(table_content))

            return '\n\n'.join(content_parts)
        except Exception as e:
            raise RuntimeError(f"读取Word文件失败: {str(e)}") from e


class MarkdownProcessor:
    """Markdown文件处理器"""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() == '.md'

    def read(self, file_path: str) -> str:
        """
        读取Markdown文件内容

        优化点：
        1. 保留原始结构信息
        2. 转换为可读的纯文本
        3. 提取标题层次
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # 先保留原始格式，进行基本清理
            cleaned = self._clean_markdown(md_content)

            # 转换为HTML再提取纯文本（保持结构）
            html = markdown.markdown(cleaned, extensions=['extra', 'toc'])

            # 简单的HTML标签清理
            text = re.sub(r'<[^>]+>', '', html)
            text = re.sub(r'\n\s*\n', '\n\n', text)

            return text.strip()
        except Exception as e:
            raise RuntimeError(f"读取Markdown文件失败: {str(e)}") from e

    def _clean_markdown(self, content: str) -> str:
        """清理Markdown内容，移除Front Matter等"""
        # 移除Jekyll风格的Front Matter
        content = re.sub(r'^---[\s\S]*?---\n', '', content)
        return content


class TextProcessor:
    """纯文本文件处理器"""

    def supports(self, file_ext: str) -> bool:
        return file_ext.lower() in ('.txt', '.text', '.log', '.csv')

    def read(self, file_path: str) -> str:
        """读取纯文本文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read().strip()
            except Exception as e:
                raise RuntimeError(f"读取文本文件失败，无法识别文件编码: {str(e)}") from e
        except Exception as e:
            raise RuntimeError(f"读取文本文件失败: {str(e)}") from e


class DocumentProcessor:
    """
    文档处理器核心类 - 统一管理各种文件格式的处理

    采用策略模式实现处理器注册和选择，支持并行处理和智能分块。

    Attributes:
        processors: 已注册的文件处理器列表
        text_splitter: 文本分块器
        max_workers: 并行处理最大线程数
    """

    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        max_workers: int = 4
    ):
        """
        初始化文档处理器

        Args:
            chunk_size: 文本分块大小，默认使用配置中的CHUNK_SIZE
            chunk_overlap: 分块重叠大小，默认使用配置中的CHUNK_OVERLAP
            max_workers: 并行处理最大线程数
        """
        self.processors: List[FileProcessor] = [
            PDFProcessor(),
            WordProcessor(),
            MarkdownProcessor(),
            TextProcessor()
        ]

        # 配置文本分块器
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        self.max_workers = max_workers

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "。", "！", "？", "，", ", ", " ", ""],
            is_separator_regex=False
        )

    def register_processor(self, processor: FileProcessor) -> None:
        """
        注册新的文件处理器

        Args:
            processor: 实现FileProcessor协议的处理器实例
        """
        self.processors.append(processor)
        print(f"已注册新处理器: {processor.__class__.__name__}")

    def get_processor(self, file_ext: str) -> Optional[FileProcessor]:
        """
        根据文件扩展名获取合适的处理器

        Args:
            file_ext: 文件扩展名（含小数点，如 '.pdf'）

        Returns:
            Optional[FileProcessor]: 找到的处理器实例，未找到返回None
        """
        for processor in self.processors:
            if processor.supports(file_ext):
                return processor
        return None

    def read_file(self, file_path: str) -> str:
        """
        读取单个文件内容（自动选择处理器）

        Args:
            file_path: 文件路径

        Returns:
            str: 提取的文本内容

        Raises:
            FileNotFoundError: 文件不存在时抛出
            ValueError: 不支持的文件类型时抛出
            RuntimeError: 文件读取失败时抛出
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()
        processor = self.get_processor(file_ext)

        if not processor:
            raise ValueError(f"不支持的文件类型: {file_ext}，"
                             f"支持类型: {self.get_supported_extensions()}")

        return processor.read(file_path)

    def split_text(self, text: str) -> List[str]:
        """
        将文本分块

        优化点：
        1. 智能分块（保留语义完整性）
        2. 记录分块位置
        3. 过滤过短的分块

        Args:
            text: 待分块的文本

        Returns:
            List[str]: 分块后的文本列表
        """
        if not text.strip():
            return []

        chunks = self.text_splitter.split_text(text)

        # 过滤过短的分块（通常是噪音）
        min_chunk_length = 50
        filtered_chunks = [
            chunk for chunk in chunks
            if len(chunk.strip()) >= min_chunk_length
        ]

        return filtered_chunks if filtered_chunks else chunks

    def process_file(self, file_path: str) -> List[Dict]:
        """
        处理单个文件，返回分块后的文档（兼容旧接口）

        Args:
            file_path: 文件路径

        Returns:
            List[Dict]: 分块文档列表，每个元素包含id、content和metadata
        """
        chunks = self.process_file_to_chunks(file_path)
        return [chunk.to_dict() for chunk in chunks]

    def process_file_to_chunks(self, file_path: str) -> List[DocumentChunk]:
        """
        处理单个文件，返回DocumentChunk对象列表

        Args:
            file_path: 文件路径

        Returns:
            List[DocumentChunk]: 分块后的文档对象列表

        Raises:
            FileNotFoundError: 文件不存在时抛出
            ValueError: 不支持的文件类型时抛出
            RuntimeError: 处理失败时抛出
        """
        filename = os.path.basename(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()

        # 读取文件内容
        text = self.read_file(file_path)

        if not text.strip():
            print(f"警告：文件 {filename} 未提取到有效内容")
            return []

        # 分块处理
        chunk_texts = self.split_text(text)

        if not chunk_texts:
            print(f"警告：文件 {filename} 分块后无有效内容")
            return []

        # 构建DocumentChunk对象
        chunks = []
        current_position = 0
        total_chunks = len(chunk_texts)

        for i, chunk_content in enumerate(chunk_texts):
            # 计算分块在原文中的位置（近似）
            start_index = current_position
            end_index = current_position + len(chunk_content)
            current_position += len(chunk_content) - self.chunk_overlap

            chunk = DocumentChunk(
                content=chunk_content,
                chunk_id=i,
                total_chunks=total_chunks,
                filename=filename,
                file_type=file_ext,
                start_index=start_index,
                end_index=end_index
            )
            chunks.append(chunk)

        print(f"成功处理文件: {filename}, 生成 {len(chunks)} 个分块")
        return chunks

    def process_files(
        self,
        file_paths: List[str],
        raise_errors: bool = False
    ) -> List[Dict]:
        """
        批量处理文件（支持并行处理）

        Args:
            file_paths: 文件路径列表
            raise_errors: 遇到错误时是否抛出异常，默认False（记录错误继续）

        Returns:
            List[Dict]: 所有成功处理的分块文档列表

        Example:
            >>> processor = DocumentProcessor()
            >>> files = ["/path/to/doc1.pdf", "/path/to/doc2.md"]
            >>> results = processor.process_files(files)
            >>> len(results)
            42
        """
        all_chunks = []

        # 使用线程池并行处理
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {
                executor.submit(self.process_file_to_chunks, path): path
                for path in file_paths
            }

            for future in as_completed(futures):
                path = futures[future]
                try:
                    chunks = future.result()
                    all_chunks.extend(chunks)
                except Exception as e:
                    error_msg = f"处理文件 {path} 失败: {str(e)}"
                    print(error_msg)
                    if raise_errors:
                        raise RuntimeError(error_msg) from e
                    continue

        # 转换为字典格式以保持兼容性
        return [chunk.to_dict() for chunk in all_chunks]

    def get_supported_extensions(self) -> List[str]:
        """
        获取当前支持的文件扩展名列表

        Returns:
            List[str]: 支持的文件扩展名列表，如 ['.pdf', '.docx', ...]
        """
        extensions = set()
        # 由于Protocol的限制，我们需要硬编码已知的扩展名
        # 实际应用中可以让每个处理器报告其支持的扩展名
        return ['.pdf', '.docx', '.doc', '.md', '.txt', '.text', '.log', '.csv']

    def validate_file(self, file_path: str) -> Dict:
        """
        验证文件是否可以正常处理

        Args:
            file_path: 文件路径

        Returns:
            Dict: 验证结果，包含success、message、file_info等字段
        """
        result = {
            "success": False,
            "filename": os.path.basename(file_path),
            "file_size": 0,
            "estimated_chunks": 0,
            "message": ""
        }

        try:
            if not os.path.exists(file_path):
                result["message"] = "文件不存在"
                return result

            # 检查文件大小
            file_size = os.path.getsize(file_path)
            result["file_size"] = file_size

            if file_size == 0:
                result["message"] = "文件为空"
                return result

            max_size = settings.MAX_FILE_SIZE
            if file_size > max_size:
                result["message"] = f"文件过大（{file_size / 1024 / 1024:.2f}MB），" \
                                    f"最大支持 {max_size / 1024 / 1024:.2f}MB"
                return result

            # 检查文件类型
            file_ext = os.path.splitext(file_path)[1].lower()
            if not self.get_processor(file_ext):
                result["message"] = f"不支持的文件类型: {file_ext}"
                return result

            # 尝试读取文件
            text = self.read_file(file_path)
            if not text.strip():
                result["message"] = "文件内容为空或无法提取文本"
                return result

            # 估算分块数量
            estimated_chunks = max(1, len(text) // self.chunk_size)
            result["estimated_chunks"] = estimated_chunks

            result["success"] = True
            result["message"] = "文件验证通过"
            result["content_length"] = len(text)

            return result

        except Exception as e:
            result["message"] = f"验证失败: {str(e)}"
            return result


# 全局便捷函数
_processor_instance: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """
    获取DocumentProcessor单例实例

    Returns:
        DocumentProcessor: 全局文档处理器实例
    """
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor()
    return _processor_instance
